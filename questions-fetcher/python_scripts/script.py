import docx
import re

def extract_questions(file_path):
    doc = docx.Document(file_path)
    questions = []
    question_pattern = re.compile(r"\(?\d+[.)]+[.]?\s*(.*?)(?=\s*[A-D]+[.\)]|$)")

    for para in doc.paragraphs:
        matches = question_pattern.findall(para.text.strip())
        for match in matches:
            question_text = match
            questions.append(question_text)

    return questions

def extract_options(file_path):
    doc = docx.Document(file_path)
    options = []
    option_pattern = re.compile(
        r"[A-D]+[.)]+[.]?\s*(.*?)(?=\s*[A-D][\.\)]|$|\s*\([A-D])", re.MULTILINE | re.DOTALL
    )

    for para in doc.paragraphs:
        matches = option_pattern.findall(para.text.strip())
        for match in matches:
            option_text = match
            options.append(option_text)

    return options

if __name__ == "__main__":
    file_path = "../data/test2.docx"
    questions = extract_questions(file_path)
    options = extract_options(file_path)
    for i, q in enumerate(questions, 1):
        print(f"Question {i}: {q}")
    for i, q in enumerate(options, 1):
        print(f"Option {i}: {q}")
