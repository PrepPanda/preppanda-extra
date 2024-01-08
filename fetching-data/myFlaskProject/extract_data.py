import docx
import re


def extract_questions(doc):
    questions = []
    question_pattern = re.compile(r"\(?\d+[.)]+[.]?\s*(.*?)(?=\s*[A-D]+[.\)]|$)")

    for para in doc.paragraphs:
        matches = question_pattern.findall(para.text.strip())
        for match in matches:
            questions_text = match
            questions.append(questions_text)

    return questions


def extract_options(doc):
    options = []
    option_pattern = re.compile(
        r"[A-D]+[.)]+[.]?\s*(.*?)(?=\s*[A-D][\.\)]|$|\s*\([A-D])", re.MULTILINE | re.DOTALL
    )

    for para in doc.paragraphs:
        matches = option_pattern.findall(para.text.strip())
        for match in matches:
            option_text = match
            options.append(option_text)

    return options


def extract_answers(doc):
    answers = []
    answer_pattern = re.compile(r"(?:Correct\s*)?(?:Ans(?:wer)?:|Answer:|Ans\.)\s*(.+)")

    for para in doc.paragraphs:
        matches = answer_pattern.findall(para.text.strip())
        for match in matches:
            answer_text = match
            answers.append(answer_text)

    return answers


# Extracts data from a .docx file and return ths json data
def extract_data(filename):
    doc = docx.Document(filename)
    questions = extract_questions(doc)
    options = extract_options(doc)
    answers = extract_answers(doc)
    data = []
    for i in range(len(questions)):
        data.append(
            {
                "question": questions[i],
                "options": [options[i], options[i + 1], options[i + 2], options[i + 3]],
                "answer": answers[i],
            }
        )
    print(data)
    return data
